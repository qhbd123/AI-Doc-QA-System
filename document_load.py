# 读取本地文档
# document_load.py作用: 将文档内容(PDF, Word, TXT, Markdown, HTML)提取成纯文本，存入知识库

from typing import List
from PyPDF2 import PdfReader
from docx import Document
import re


# 可选：用于解析 HTML 的库（如果未安装，请运行 pip install beautifulsoup4）
try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True  # 将变量 HAS_BS4 设置为 True，表示 BeautifulSoup 可用
except ImportError:
    HAS_BS4 = False
    print("[doc_load] 警告: 未安装 BeautifulSoup，HTML 文件将仅提取文本（不处理标签）")

# ---------------------- 1. 读取不同格式文档 ----------------------
def load_txt(file_path: str) -> str:
    """读取 TXT 文档，兼容多种编码"""
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ValueError(f"无法识别文件编码：{file_path}")


def load_pdf(file_path: str) -> str:
    """读取 PDF 文档，处理空页与异常"""
    text = ""
    try:
        # 创建 PDF 读取器对象，使用 PyPDF2 库中的 PdfReader 类打开指定路径的 PDF 文件
        reader = PdfReader(file_path)
        # 遍历每一页
        for page in reader.pages:
            # 提取当前页的文本
            page_text = page.extract_text()
            # 过滤空白页并累积文本
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        raise RuntimeError(f"读取PDF失败：{str(e)}")
    return text


def load_docx(file_path: str) -> str:
    """读取 Word（.docx）文档"""
    # 创建 Word 文档对象，使用 python-docx 库中的 Document 类打开指定路径的 .docx 文件
    doc = Document(file_path)
    # 提取并拼接段落文本，文档中所有非空段落的文本，段落之间用换行符隔开
    # doc.paragraphs：文档对象的一个属性，返回文档中所有段落的可迭代序列
    # para.text 返回段落的文本内容（字符串），.strip() 去除首尾空白字符
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])


def load_md(file_path: str) -> str:
    """读取 Markdown 文档，去除 Markdown 语法，保留纯文本"""
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    # 移除代码块:删除 Markdown 文本中的所有代码块（被三个反引号包裹的内容），将其替换为空字符串
    # r"```.*?```" 匹配以 ``` 开始、.*?：匹配任意字符（.）零次或多次、以 ``` 结束的最短内容
    # re.DOTALL 使得正则中的 . 能够匹配包括换行符在内的任何字符
    content = re.sub(r"```.*?```", "", content, flags=re.DOTALL)
    # 移除行内代码标记，去掉反引号，保留内部代码文本
    content = re.sub(r"`([^`]+)`", r"\1", content)
    # 移除标题标记 # ## ### 等
    content = re.sub(r"^#{1,6}\s+", "", content, flags=re.MULTILINE)
    # 移除粗体/斜体标记 **text** 或 __text__ 或 *text* 或 _text_
    content = re.sub(r"\*\*(.*?)\*\*", r"\1", content)
    content = re.sub(r"__(.*?)__", r"\1", content)
    content = re.sub(r"\*(.*?)\*", r"\1", content)
    content = re.sub(r"_(.*?)_", r"\1", content)
    # 移除链接 [text](url) -> text，将 Markdown 链接替换为 显示文本，只保留文字部分
    content = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", content)
    # 移除图片 ![alt](url)
    content = re.sub(r"!\[[^\]]*\]\([^\)]+\)", "", content)
    # 移除水平线 --- 或 ***
    content = re.sub(r"^[-*]{3,}\s*$", "", content, flags=re.MULTILINE)
    # 移除列表标记 - 或 * 或 数字.
    content = re.sub(r"^[\s]*[-*+]\s+", "", content, flags=re.MULTILINE)
    content = re.sub(r"^\s*\d+\.\s+", "", content, flags=re.MULTILINE)
    # 去除处理后的字符串首尾的空白字符
    return content.strip()


def load_html(file_path: str) -> str:
    """读取 HTML 文档，提取纯文本"""
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    if HAS_BS4:
        # 使用 html.parser 解析器将 HTML 字符串转换为 BeautifulSoup 对象，可以方便地操作 DOM 树
        soup = BeautifulSoup(content, "html.parser")
        # 移除脚本和样式
        for script in soup(["script", "style"]):
            script.decompose()
        # get_text() 方法遍历 DOM 树，提取所有可见文本，并用空格分隔不同元素的内容
        text = soup.get_text()
        # 清理多余空白
        # text.splitlines()：将文本按换行符分割成行列表
        # (line.strip()去除每行首尾空白字符
        lines = (line.strip() for line in text.splitlines())
        # "\n".join(...)：用换行符将非空行重新连接，使文本保留段落分隔但无多余空白
        text = "\n".join(line for line in lines if line)
        return text
    else:
        # 降级方案：简单正则去除标签
        # 用空字符串替换所有形如 <tag> 或 </tag> 的 HTML 标签
        text = re.sub(r"<[^>]+>", "", content)
        # 将连续的空白字符替换为单个空格，避免因删除标签后留下大量空白
        text = re.sub(r"\s+", " ", text).strip()
        return text


# ---------------------- 2. 文本清洗 ---------------------
def clean_text(text: str) -> str:
    """清理冗余空白，但保留换行作为段落分隔"""
    # 统一换行符，将不同操作系统下的换行符统一为 Unix/Linux 风格的 \n
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # 合并多个连续换行为两个换行（保留段落边界）
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    # 将连续的空格和制表符压缩为一个空格，但不影响换行符
    text = re.sub(r"[ \t]+", " ", text)
    # 移除字符串开头和结尾的空白字符
    return text.strip()


# ---------------------- 3. 语义化文档拆分 ----------------------
def _split_long_paragraph(para: str, max_size: int, delim_pattern) -> List[str]:
    """
    当段落过长时，按句子边界将其拆分成多个较小的片段，同时尽量保持句子的完整性，避免在单词中间截断
    """
    # 提取句子（保留分隔符）
    sentences = []
    last_end = 0
    # 在 para 中查找所有匹配 delim_pattern 的位置
    for match in delim_pattern.finditer(para):
        # 获取当前匹配的起始和结束索引
        start, end = match.span()
        if start > last_end:
            sentences.append(para[last_end:end].strip())
        last_end = end
    if last_end < len(para):
        sentences.append(para[last_end:].strip())
    # 过滤空句子
    sentences = [s for s in sentences if s]
    if not sentences:
        return []

    # 合并句子为 chunk（不超过 max_size）
    chunks = []
    current = ""
    for s in sentences:
        # 每个句子末尾加换行（保留分隔）
        if len(current) + len(s) + 2 <= max_size:
            current += s + "\n"
        else:
            if current:
                chunks.append(current.strip())
            current = s + "\n"
    if current:
        chunks.append(current.strip())
    return chunks


def split_text(text: str, max_chunk_size: int = 500) -> List[str]:
    """
    RAG 系统中文档分块的核心函数
    将长文本按「段落 → 句子」的层次进行智能拆分，保证每个块长度不超过 max_chunk_size，同时尽量保持句子完整，避免在单词中间截断
    """
    if not text or not text.strip():  # 文本为空或含空白符号
        return []

    # 使用正则表达式 \n\s*\n+ 分割文本
    # 对每个分割出的段落片段去除首尾空白，并过滤掉空字符串，得到段落列表 paragraphs
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", text) if p.strip()]
    chunks = []
    current_chunk = ""

    # 句子边界分隔符正则（中文/英文句号、感叹号、问号）
    sentence_delimiters = re.compile(r'[。！？!?\.]')

    for p in paragraphs:
        # 尝试合并段落，不超过最大长度
        if len(current_chunk) + len(p) + 2 <= max_chunk_size:
            current_chunk += p + "\n"
        else:
            # 先把当前累积的 chunk 存起来
            if current_chunk:
                chunks.append(current_chunk.strip())
            # 把过长的段落按句子拆分
            para_chunks = _split_long_paragraph(p, max_chunk_size, sentence_delimiters)
            if para_chunks:
                chunks.extend(para_chunks)
                current_chunk = ""
            else:
                # 如果拆分失败（无句子），降级为整体添加（限制长度）
                if len(p) > max_chunk_size:
                    chunks.append(p[:max_chunk_size])
                else:
                    current_chunk = p + "\n"
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks


# ---------------------- 4. 统一入口函数 ----------------------
def load_document(file_path: str, max_chunk_size: int = 500) -> List[str]:
    """
    统一加载文档入口，自动识别格式 → 清洗 → 拆分
    支持格式: txt, pdf, docx, md, html
    :param file_path: 文档路径
    :param max_chunk_size: 片段最大长度
    :return: 拆分后的文本片段列表
    """
    print(f"[doc_load] 加载文档: {file_path}")
    ext = file_path.lower().split(".")[-1]
    if ext == "txt":
        text = load_txt(file_path)
    elif ext == "pdf":
        text = load_pdf(file_path)
    elif ext == "docx":
        text = load_docx(file_path)
    elif ext == "md":
        text = load_md(file_path)
    elif ext == "html" or ext == "htm":
        text = load_html(file_path)
    else:
        raise ValueError(f"不支持的文件格式：{ext}")

    text = clean_text(text)
    print(f"[doc_load] 清洗后文本长度: {len(text)} 字符")
    chunks = split_text(text, max_chunk_size)
    print(f"[doc_load] 拆分完成: {len(chunks)} 个片段")
    return chunks
